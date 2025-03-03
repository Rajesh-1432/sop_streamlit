import io
import os
import json
import fitz  # PyMuPDF
import zipfile
import tempfile
import streamlit as st
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from dotenv import load_dotenv
import base64
from openai import OpenAI

load_dotenv()
# Initialize OpenAI client (you'll need to set OPENAI_API_KEY in your environment variables)
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# Define paths for reference files
REFERENCE_IMAGE_PATH = "References\\ref_image.png"
REFERENCE_TEXT_PATH = "References\\ref_output.txt"

def encode_image_to_base64(image):
    """
    Convert PIL Image to base64 for OpenAI API.
    """
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    img_str = base64.b64encode(buffered.getvalue()).decode('utf-8')
    return img_str


def analyze_process_flow_image(image, reference_image_path=None, reference_text_path=None):
    """
    Analyze a process flow image using OpenAI's vision model.
    """
    try:
        # Convert PIL image to base64
        image_base64 = encode_image_to_base64(image)
        
        # Prepare system message and content
        content = []
        system_message = "You are an expert at analyzing process flow diagrams and converting them into detailed text descriptions."
        
        # Add reference image if provided
        if reference_image_path and os.path.exists(reference_image_path):
            with open(reference_image_path, 'rb') as ref_file:
                ref_image_bytes = ref_file.read()
                ref_image = Image.open(io.BytesIO(ref_image_bytes))
                ref_image_base64 = encode_image_to_base64(ref_image)
                content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{ref_image_base64}"
                    }
                })
                content.append({
                    "type": "text",
                    "text": "Additional Context: Here is a reference image for additional context:"
                })
        
        # Add reference text if provided
        if reference_text_path and os.path.exists(reference_text_path):
            with open(reference_text_path, 'r', encoding='utf-8') as file:
                reference_text = file.read()
                content.append({
                    "type": "text",
                    "text": f"Output Format:\n{reference_text}"
                })

        # Add analysis instructions and image
        content.append({
            "type": "text", 
            "text": """Analyze this process flow diagram. 
            Describe the steps in detail in such a way that it is shown in the "Output Format".
            Use the Output Format given above for generating a response.
            Generate an Objective and also the Purpose for the processflow(Image) in 3-4 sentences.
            The steps should be ordered in such a way that the processflow image is there.
            Consider all possible flows if there are multiple options after a step create a and b for those steps.
            So understand the pattern and generate the response based on the "Output Format".
            Fix the output format and dont deviate from it.
            Consider all the boxes in the Image as Step and Create sub steps for each step similar to that of Output Reference.
            In the details step try to add as many steps as possible for each substep.
            Do not consider reference text as the input it is just for understanding the output 
            Do not use the reference text in the output
            IMPORTANT: Provide the response in valid JSON format with the following structure:
            {
              "title": "...",
              "Objective": "...",
              "purpose": "...",
              "steps": [
                {
                  "step": "...",
                  "role": "...",
                  "activities": [
                    {
                      "task": "...",
                      "details": [
                        "...",
                        "..."
                      ]
                    }
                  ]
                }
              ]
            }"""
        })
        
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/png;base64,{image_base64}"
            }
        })

        # Generate response using OpenAI
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": content}
            ],
        )
        
        response_text = response.choices[0].message.content.strip()

        try:
            # First attempt: Try to parse the entire response as JSON
            return json.loads(response_text)
        except json.JSONDecodeError:
            try:
                # Second attempt: Try to find JSON content between markers
                start_marker = "```json"
                end_marker = "```"
                if start_marker in response_text and end_marker in response_text:
                    json_content = response_text.split(start_marker)[1].split(end_marker)[0].strip()
                    return json.loads(json_content)
                else:
                    raise Exception("No valid JSON content found in response")
            except Exception as e:
                print(f"Error parsing JSON content: {e}")
                print("Raw response:", response_text)
                return None
    
    except Exception as e:
        print(f"Error analyzing image with OpenAI: {e}")
        return None

def extract_images_from_pdf(pdf_path):
    """
    Extract images from a PDF file using PyMuPDF.
    """
    images = []
    try:
        pdf_document = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            image_list = page.get_images(full=True)
            
            for img_index, img_info in enumerate(image_list):
                try:
                    base_image = pdf_document.extract_image(img_info[0])
                    image_data = base_image["image"]
                    pil_image = Image.open(io.BytesIO(image_data))
                    images.append(pil_image)
                
                except Exception as img_error:
                    print(f"Error extracting image {img_index} from page {page_num}: {img_error}")
        
        pdf_document.close()
    
    except Exception as e:
        print(f"Error processing PDF: {e}")
    
    return images

def save_image_for_doc(image, output_dir):
    """
    Save a PIL Image to a temporary file for document insertion.
    """
    temp_image_path = os.path.join(output_dir, 'process_flow_map.png')
    image.save(temp_image_path)
    return temp_image_path

def add_table_borders(table):
    tbl = table._element
    tbl_pr = tbl.xpath(".//w:tblPr")[0]
    tbl_borders = parse_xml(
        """
        <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tblBorders>
        """
    )
    tbl_pr.append(tbl_borders)

def create_docx_from_analysis(analysis_json, output_path, process_flow_image):
    """
    Create a DOCX file from the analysis JSON with proper formatting.
    """
    doc = Document()
    
    # Define styles
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.size = Pt(26)
    title_font.bold = True
    title_font.color.rgb = RGBColor(16,129,242)

    title_style = doc.styles.add_style('Custom', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.size = Pt(20)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0,0,0)
    
    heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_font = heading1_style.font
    heading1_font.size = Pt(14)
    heading1_font.bold = True
    heading1_font.color.rgb=RGBColor(0,0,0)
    
    heading2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_font = heading2_style.font
    heading2_font.size = Pt(12)
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(0,0,0)

    # Add title
    if 'title' in analysis_json:
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        title_para = doc.add_paragraph(analysis_json['title'])
        title_para.style = 'CustomTitle'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para = doc.add_paragraph("Standard Operating Procedure")
        title_para.style = 'Custom'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        section = doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph("Document History", style='CustomHeading1')
    doc.add_paragraph('Document Location', style='CustomHeading2')
    doc.add_paragraph('This is an on-line document. Paper copies are valid only on the day they are printed. Refer to the Genpact approver for location where last version of the document is stored or if you are in any doubt about the accuracy of this document')
    # Add Process Flow Map heading
    doc.add_paragraph('Document Creation', style='CustomHeading2')
    table = doc.add_table(rows=2, cols=3)
    add_table_borders(table)  # Add borders to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Creation Date'
    hdr_cells[1].text = 'Genpact approval by'
    hdr_cells[2].text = 'Customer approval by'

    doc.add_paragraph('Revision History', style='CustomHeading2')
    table = doc.add_table(rows=4, cols=5)
    add_table_borders(table)  # Add borders to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Revision Date'
    hdr_cells[1].text = 'Version Number'
    hdr_cells[2].text = 'Change Reason'
    hdr_cells[3].text = 'Pages Changed'
    hdr_cells[4].text = 'Approval By'
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph('Table of Contents', style='CustomHeading1')
    toc_items = [
        ' Overview',
        '   Purpose and Scope',
        '   Definitions',
        '   System of Engagement',
        '   Roles and Responsibilities',
        ' Process Narrative',
        '   COPIS',
        '   Process Map/Flowchart',
        ' Detailed Process Steps',
        ' Process Exceptions Handling',
        ' Compliance Control',
        ' Escalation Process',
        ' Process SLAs',
        ' Related Documents',
        ' Sign Off'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph('Overview', style='CustomHeading1')
    doc.add_paragraph('     Purpose and Scope', style='CustomHeading2')
    doc.add_paragraph(analysis_json.get('Objective', 'N/A'))
    doc.add_paragraph('     Definitions', style='CustomHeading2')
    doc.add_paragraph('         Acronyms', style='CustomHeading2')
    table = doc.add_table(rows=4, cols=2)
    add_table_borders(table)  # Add borders to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Abbreviation:'
    hdr_cells[1].text = 'Long Form:'
    doc.add_paragraph('         Definitions', style='CustomHeading2')
    table = doc.add_table(rows=4, cols=2)
    add_table_borders(table)  # Add borders to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Term:'
    hdr_cells[1].text = 'Definition:'
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph("System of Engagement", style='CustomHeading1')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Roles and Responsibilities in performing this activity', style='CustomHeading1')
    table = doc.add_table(rows=4, cols=2)
    add_table_borders(table)  # Add borders to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role:'
    hdr_cells[1].text = 'Responsibility:'
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph("Process Narrative", style='CustomHeading1') 
    doc.add_paragraph(analysis_json.get('purpose', 'N/A'))   
    doc.add_paragraph("Process Flow Map", style='CustomHeading1')
    doc.add_paragraph()

    # Save and insert the process flow image
    output_dir = os.path.dirname(output_path)
    image_path = save_image_for_doc(process_flow_image, output_dir)
    
    # Insert the image with specified width
    doc.add_picture(image_path, width=Inches(6.0))
    
    # Add spacing after image
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph("Detailed Process Steps", style='CustomHeading1') 
    # Process other top-level keys
    if 'steps' in analysis_json:
        steps = analysis_json['steps']
        
        if isinstance(steps, str):
            doc.add_paragraph(steps)
        elif isinstance(steps, list):
            for item in steps:
                step_heading = f"Step {item.get('step', 'N/A')}: {item.get('role', 'N/A')}"
                doc.add_paragraph(step_heading, style='CustomHeading2')
                
                if 'activities' in item and isinstance(item['activities'], list):
                    for activity in item['activities']:
                        task_para = doc.add_paragraph()
                        task_para.add_run(f"Task: {activity.get('task', 'N/A')}").bold = True
                        
                        if 'details' in activity and isinstance(activity['details'], list):
                            for detail in activity['details']:
                                doc.add_paragraph(f"{detail}", style='List Bullet')
                        doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_paragraph("Process Exception Handling", style='CustomHeading1') 
    doc.add_paragraph()
    doc.add_paragraph("Compliance control", style='CustomHeading1') 
    doc.add_paragraph()
    doc.add_paragraph("Escalation Process", style='CustomHeading1') 
    table = doc.add_table(rows=3, cols=4)
    add_table_borders(table)  # Add borders to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Escalation Level:'
    hdr_cells[1].text = 'Name of Contact'
    hdr_cells[2].text = 'Title'
    hdr_cells[3].text = 'Email'
    doc.add_paragraph("Process SLAs", style='CustomHeading1')
    table = doc.add_table(rows=4, cols=5)
    add_table_borders(table)  # Add borders to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Indicator:'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Operational Definition'
    hdr_cells[3].text = 'Target'
    hdr_cells[4].text = 'Minimum Level'
    doc.add_paragraph("Related Documents", style='CustomHeading1')
    doc.add_paragraph() 
    doc.add_paragraph("Sign Off", style='CustomHeading1') 

    # Save the document
    doc.save(output_path)

    # Clean up temporary image file
    try:
        os.remove(image_path)
    except:
        pass

def process_single_pdf(pdf_data):
    """
    Process a single PDF file and return the path to the generated DOCX.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        # Save PDF to temporary file
        temp_pdf = os.path.join(temp_dir, "temp.pdf")
        with open(temp_pdf, "wb") as f:
            f.write(pdf_data)
        
        # Extract and process images
        images = extract_images_from_pdf(temp_pdf)
        
        if not images:
            return None
        
        # Process first image (assuming one process flow per PDF)
        image = images[0]
        analysis = analyze_process_flow_image(
            image, 
            reference_image_path=REFERENCE_IMAGE_PATH,
            reference_text_path=REFERENCE_TEXT_PATH
        )
        
        if analysis:
            if isinstance(analysis, str):
                analysis_json = json.loads(analysis)
            else:
                analysis_json = analysis
            
            # Create output document
            output_path = os.path.join(temp_dir, "output.docx")
            create_docx_from_analysis(analysis_json, output_path, image)
            
            # Read the generated file
            with open(output_path, "rb") as f:
                return f.read()
    
    return None

def process_zip_file(zip_data):
    """
    Process multiple PDFs from a zip file and return a zip file containing all outputs.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        # Extract zip contents
        input_zip_path = os.path.join(temp_dir, "input.zip")
        with open(input_zip_path, "wb") as f:
            f.write(zip_data)
        
        # Create directory for extracted files
        extract_dir = os.path.join(temp_dir, "extracted")
        os.makedirs(extract_dir)
        
        # Extract files
        with zipfile.ZipFile(input_zip_path, 'r') as zip_ref:
            zip_ref.extractall(extract_dir)
        
        # Create directory for outputs
        output_dir = os.path.join(temp_dir, "outputs")
        os.makedirs(output_dir)
        
        # Process each PDF file
        for root, _, files in os.walk(extract_dir):
            for filename in files:
                if filename.lower().endswith('.pdf'):
                    pdf_path = os.path.join(root, filename)
                    with open(pdf_path, "rb") as f:
                        pdf_data = f.read()
                    
                    output_data = process_single_pdf(pdf_data)
                    if output_data:
                        output_name = os.path.splitext(filename)[0] + ".docx"
                        output_path = os.path.join(output_dir, output_name)
                        with open(output_path, "wb") as f:
                            f.write(output_data)
        
        # Create output zip file
        output_zip_path = os.path.join(temp_dir, "output.zip")
        with zipfile.ZipFile(output_zip_path, 'w') as zipf:
            for root, _, files in os.walk(output_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, output_dir)
                    zipf.write(file_path, arcname)
        
        # Read the output zip file
        with open(output_zip_path, "rb") as f:
            return f.read()
    
    return None

def main():
    st.title("Process Flow Analysis")
    
    # File uploader for PDFs
    uploaded_files = st.file_uploader(
        "Upload PDF files or a ZIP containing PDFs", 
        type=['pdf', 'zip'],
        accept_multiple_files=True
    )
    
    if uploaded_files:
        if st.button("Process Files"):
            with st.spinner("Processing files..."):
                if len(uploaded_files) == 1 and uploaded_files[0].name.lower().endswith('.zip'):
                    # Process zip file
                    output_data = process_zip_file(uploaded_files[0].getvalue())
                    if output_data:
                        st.download_button(
                            "Download Results",
                            output_data,
                            "Generated_SOP's.zip",
                            "application/zip"
                        )
                    else:
                        st.error("Error processing zip file")
                
                elif len(uploaded_files) == 1:
                    # Process single PDF
                    output_data = process_single_pdf(uploaded_files[0].getvalue())
                    if output_data:
                        output_name = os.path.splitext(uploaded_files[0].name)[0] + ".docx"
                        st.download_button(
                            "Download Result",
                            output_data,
                            output_name,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("Error processing PDF file")
                
                else:
                    # Process multiple files
                    with tempfile.TemporaryDirectory() as temp_dir:
                        # Create zip file containing all PDFs
                        input_zip_path = os.path.join(temp_dir, "input.zip")
                        with zipfile.ZipFile(input_zip_path, 'w') as zipf:
                            for uploaded_file in uploaded_files:
                                zipf.writestr(uploaded_file.name, uploaded_file.getvalue())
                        
                        # Process the zip file
                        with open(input_zip_path, "rb") as f:
                            output_data = process_zip_file(f.read())
                        
                        if output_data:
                            st.download_button(
                                "Download Results",
                                output_data,
                                "Generated_SOP's.zip",
                                "application/zip"
                            )
                        else:
                            st.error("Error processing files")

if __name__ == "__main__":
    main()